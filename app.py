import os
import json
import hashlib
import requests
import csv
import io
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
from datetime import datetime
from flask import Flask, request, jsonify, render_template_string, Response
from flask_cors import CORS
from werkzeug.utils import secure_filename
from functools import wraps
import PyPDF2
import pdfplumber
import docx
from openai import OpenAI
from dotenv import load_dotenv
import re
from openpyxl import Workbook

load_dotenv()

app = Flask(__name__)
CORS(app)

US_STATES = [
    "alabama","alaska","arizona","arkansas","california","colorado","connecticut","delaware","florida","georgia",
    "hawaii","idaho","illinois","indiana","iowa","kansas","kentucky","louisiana","maine","maryland","massachusetts",
    "michigan","minnesota","mississippi","missouri","montana","nebraska","nevada","new hampshire","new jersey",
    "new mexico","new york","north carolina","north dakota","ohio","oklahoma","oregon","pennsylvania","rhode island",
    "south carolina","south dakota","tennessee","texas","utah","vermont","virginia","washington","west virginia",
    "wisconsin","wyoming","district of columbia"
]

def classify_location(address_text: str, phone: str) -> str:
    t = address_text.lower()
    
    has_state = any(state in t for state in US_STATES)
    has_zip = bool(re.search(r"\b\d{5}(?:-\d{4})?\b", address_text))
    
    has_us_phone = phone.strip().startswith("+1") or re.match(r"^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$", phone)
    
    if has_state or has_zip or has_us_phone:
        return "Onshore"
    return "Offshore"

UPLOAD_FOLDER = 'resumes'
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
ALLOWED_MIME_TYPES = {
    'application/pdf': 'pdf',
    'application/msword': 'doc', 
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx'
}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
CANDIDATES_DB_FILE = 'candidates_db.json'
SCORING_SETTINGS_FILE = 'scoring_settings.json'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
ZOHO_FLOW_WEBHOOK = os.getenv('ZOHO_FLOW_WEBHOOK', '')
ADMIN_USERNAME = os.getenv('ADMIN_USERNAME', 'admin')
ADMIN_PASSWORD = os.getenv('ADMIN_PASSWORD', 'secure123')

def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        if not auth or not (auth.username == ADMIN_USERNAME and auth.password == ADMIN_PASSWORD):
            # Send WWW-Authenticate header to trigger browser login prompt
            response = jsonify({'error': 'Authentication required'})
            response.headers['WWW-Authenticate'] = 'Basic realm="Admin Access"'
            return response, 401
        return f(*args, **kwargs)
    return decorated

def validate_file(file):
    if file.filename == '':
        return False, "No file selected"
    
    # Check file size
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    
    if size > MAX_FILE_SIZE:
        return False, f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
    
    if size == 0:
        return False, "File is empty"
    
    # Check file extension
    if not allowed_file(file.filename):
        return False, "Invalid file type. Only PDF and Word documents allowed"
    
    # Advanced MIME type checking if magic is available
    if MAGIC_AVAILABLE:
        file_content = file.read(1024)  # Read first 1KB
        file.seek(0)
        
        try:
            mime_type = magic.from_buffer(file_content, mime=True)
            if mime_type not in ALLOWED_MIME_TYPES:
                return False, f"Invalid file type. Only PDF and Word documents allowed. Got: {mime_type}"
            
            # Verify extension matches MIME type
            expected_ext = ALLOWED_MIME_TYPES[mime_type]
            actual_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
            
            if actual_ext != expected_ext:
                return False, f"File extension doesn't match content type"
                
        except Exception as e:
            # Fall back to extension checking if magic fails
            print(f"Magic validation failed, using extension check: {str(e)}")
    else:
        # Basic file signature checking without magic
        file_content = file.read(8)  # Read first 8 bytes
        file.seek(0)
        
        # Check basic file signatures
        if file_content.startswith(b'%PDF-'):
            # PDF file
            if not file.filename.lower().endswith('.pdf'):
                return False, "File appears to be PDF but has wrong extension"
        elif file_content.startswith(b'PK\x03\x04'):
            # ZIP-based format (likely DOCX)
            if not file.filename.lower().endswith(('.docx', '.doc')):
                return False, "File appears to be Office document but has wrong extension"
    
    return True, "Valid file"

def get_scoring_settings():
    if os.path.exists(SCORING_SETTINGS_FILE):
        with open(SCORING_SETTINGS_FILE, 'r') as f:
            return json.load(f)
    return {
        "version": "1.0",
        "scoring_rules": {"years_5plus": 30, "certifications": 20, "qa_training": 10, "lsp_experience": 10},
        "tier_thresholds": {"tier_1_min": 80, "tier_2_min": 60},
        "known_lsps": ["LanguageLine", "TransPerfect", "Propio", "Lionbridge"],
        "remote_keywords": ["VRI", "OPI", "Remote Interpreting", "Phone Interpreting", "Video Interpreting"]
    }

def save_scoring_settings(settings):
    settings['last_updated'] = datetime.now().isoformat()
    with open(SCORING_SETTINGS_FILE, 'w') as f:
        json.dump(settings, f, indent=2)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except (Exception) as e:
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as fallback_error:
            raise Exception(f"Failed to extract PDF text: {str(e)}, fallback error: {str(fallback_error)}")
    return text.strip()

def convert_docx_to_pdf(docx_path):
    import re
    from docx2pdf import convert
    
    converted_dir = os.path.join(UPLOAD_FOLDER, 'converted_pdfs')
    os.makedirs(converted_dir, exist_ok=True)
    
    pdf_path = os.path.join(converted_dir, os.path.basename(docx_path).replace('.docx', '.pdf').replace('.doc', '.pdf'))
    
    try:
        convert(docx_path, pdf_path)
        return pdf_path
    except Exception as e:
        print(f"Error converting DOCX to PDF: {str(e)}")
        return None

def clean_and_fix_text(text):
    import re
    text = re.sub(r'\s+@\s+', '@', text)
    text = re.sub(r'(\w)\s*\.\s*(\w)', r'\1.\2', text)
    return text.strip()

def extract_text_from_docx(filepath):
    text = ""
    
    try:
        import mammoth
        with open(filepath, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text = result.value
            if text and len(text.strip()) > 50:
                return clean_and_fix_text(text)
    except Exception as e:
        pass
    
    pdf_path = convert_docx_to_pdf(filepath)
    if pdf_path and os.path.exists(pdf_path):
        text = extract_text_from_pdf(pdf_path)
        if text and len(text.strip()) > 50:
            return clean_and_fix_text(text)
    
    try:
        doc = docx.Document(filepath)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        if doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
        
        if text and len(text.strip()) > 10:
            return clean_and_fix_text(text)
    except Exception as e:
        pass
    
    try:
        import subprocess
        result = subprocess.run(['textutil', '-convert', 'txt', '-stdout', filepath], 
                              capture_output=True, text=True, timeout=10)
        if result.stdout and len(result.stdout.strip()) > 10:
            return clean_and_fix_text(result.stdout)
    except Exception as e:
        print(f"Error using textutil: {str(e)}")
        pass
    
    return text if text else ""

def extract_text(filepath, file_extension):
    if file_extension == 'pdf':
        text = extract_text_from_pdf(filepath)
        return clean_and_fix_text(text)
    elif file_extension in ['doc', 'docx']:
        return extract_text_from_docx(filepath)
    return ""

def get_candidates_db():
    if os.path.exists(CANDIDATES_DB_FILE):
        with open(CANDIDATES_DB_FILE, 'r') as f:
            return json.load(f)
    return {}

def save_candidate(candidate_id, data):
    db = get_candidates_db()
    db[candidate_id] = data
    with open(CANDIDATES_DB_FILE, 'w') as f:
        json.dump(db, f, indent=2)

def generate_identifier(text):
    import re
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    if email_match:
        return email_match.group(0).lower()
    return None

def validate_parsed_data(data):
    required_fields = ['name', 'email', 'tier_level', 'tier_score', 'qualify']
    for field in required_fields:
        if field not in data:
            return False
        if field == 'tier_score':
            if data[field] is None or not isinstance(data[field], (int, float)):
                return False
        elif not data[field]:
            return False
    return True

def parse_resume_with_openai(text):
    settings = get_scoring_settings()
    scoring_rules = settings['scoring_rules']
    tier_thresholds = settings['tier_thresholds']
    known_lsps = ', '.join(settings['known_lsps'])
    remote_keywords = ', '.join(settings['remote_keywords'])
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an AI assistant that processes resumes and outputs structured candidate data. Always return a complete JSON object with all fields, even if information is missing or not relevant. If a field is missing in the resume, use \"\" for strings, [] for arrays, 0 for numbers, and false for booleans. Never omit keys. If candidate has no interpreting experience, classify them as Tier 3 with tier_score: 0."},
                {"role": "user", "content": f"""Process this resume and return the following JSON with ALL fields filled.

{{
  "name": "",
  "email": "",
  "primary_language": "",
  "other_spoken_languages": [],
  "service_location": "Onshore | Offshore | Unknown",
  "mobile": "",
  "remote_experience": false,
  "tier_level": "Tier 1 | Tier 2 | Tier 3",
  "tier_score": 0,
  "education": "None | Associate's Degree | Bachelor's Degree | Currently Enrolled - Graduate | Currently Enrolled - Undergraduate | Doctorate (Ph.D.) | Graduate | High School Diploma | Master's Degree | No Formal Education | Post Graduate | Professional Degree | Some College (No Degree) | Undergraduate",
  "qualify": "Yes - Qualified | Not Qualified",
  "role_relevance": "Interpreter | Translator | Not Relevant",
  "training_needed": false,
  "processing_notes": "",
  "certifications": [],
  "skills": [],
  "experience": [
    {{
      "company": "",
      "position": "",
      "duration": "",
      "description": ""
    }}
  ],
  "address": {{
    "street": "",
    "city": "",
    "state": "",
    "zip_code": "",
    "country": ""
  }}
}}

Rules for Classification:

Role Relevance:
- "Interpreter" if interpreting experience is found (remote or on-site)
- "Translator" if only translation/subtitling/QA found
- "Not Relevant" if resume belongs to another domain (HR, IT, etc.)

Remote Experience:
- If resume mentions {remote_keywords} → "remote_experience": true
- Otherwise → "remote_experience": false

Tier Assignment:
- Tier 1: Has remote interpreting (OPI/VRI) experience + Score ≥{tier_thresholds['tier_1_min']}
- Tier 2: On-site interpreting only (court, hospital, community), no remote interpreting → even if Score ≥{tier_thresholds['tier_1_min']}, cap Tier at 2
- Tier 3: No interpreting experience → Score = 0

Scoring (0–100):
- +{scoring_rules['years_5plus']} if >5 years interpreting experience
- +{scoring_rules['certifications']} for certifications
- +{scoring_rules['qa_training']} for QA/Training experience
- +{scoring_rules['lsp_experience']} if worked for known LSPs (e.g. {known_lsps})
- Base score should never override Tier logic

Qualification:
- "Yes - Qualified" if interpreting experience exists (remote or on-site)
- "Not Qualified" if no interpreting experience

Training Needed:
- "training_needed": true for Tier 2 and Tier 3 (they require training before remote assignments)
- "training_needed": false for Tier 1

Processing Notes:
- Always provide a short explanation (e.g., "On-site court interpreter only, needs VRI training" or "No interpreting experience, Tier 3 by default")

Service Location:
- IMPORTANT: Leave as "Unknown" in your JSON response
- The system will automatically determine location based on address and phone number

Education: Must match one of the listed categories exactly.

Examples:
- Court interpreter, 10 years on-site only → Tier 2, Score ~{scoring_rules['years_5plus'] + scoring_rules['certifications']}, training_needed: true, processing_notes: "On-site court interpreter only, needs VRI training"
- Hospital interpreter with OPI/VRI → Tier 1, Score {scoring_rules['years_5plus'] + scoring_rules['certifications'] + scoring_rules['lsp_experience']}, training_needed: false, processing_notes: "Experienced remote interpreter"
- Translator, no interpreting → Tier 3, Score 0, training_needed: true, processing_notes: "Translator only, no interpreting experience"
- HR professional → Tier 3, Score 0, role_relevance: "Not Relevant", training_needed: true, processing_notes: "No interpreting experience, Tier 3 by default"

Resume text:
{text}

Return only valid JSON, no additional text."""}
            ],
            temperature=0.3
        )
        
        result = response.choices[0].message.content.strip()
        if result.startswith('```json'):
            result = result[7:]
        if result.startswith('```'):
            result = result[3:]
        if result.endswith('```'):
            result = result[:-3]
        
        parsed = json.loads(result.strip())
        return parsed, None
    except json.JSONDecodeError as e:
        return None, f"JSON parsing error: {str(e)} | Raw response: {result[:200] if 'result' in locals() else 'N/A'}"
    except Exception as e:
        return None, f"Error: {str(e)}"

def send_to_zoho_flow(candidate_data):
    if not ZOHO_FLOW_WEBHOOK:
        return False, "Zoho Flow webhook not configured"
    
    try:
        response = requests.post(ZOHO_FLOW_WEBHOOK, json=candidate_data, timeout=10)
        if response.status_code == 200:
            return True, "Sent to Zoho Flow successfully"
        else:
            return False, f"Zoho Flow returned status {response.status_code}"
    except Exception as e:
        return False, str(e)

@app.route('/upload', methods=['POST'])
def upload_resume():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    # Use enhanced file validation
    is_valid, error_message = validate_file(file)
    if not is_valid:
        return jsonify({'error': error_message}), 400
    
    filename = secure_filename(file.filename)
    file_extension = filename.rsplit('.', 1)[1].lower()
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    
    file.save(filepath)
    
    text = extract_text(filepath, file_extension)
    
    if not text:
        return jsonify({'error': 'Could not extract text from file', 'status': 'failed'}), 400
    
    identifier = generate_identifier(text)
    
    if not identifier:
        identifier = hashlib.md5(filename.encode()).hexdigest()
    
    db = get_candidates_db()
    
    if identifier in db:
        existing = db[identifier]
        if existing['status'] == 'processed':
            return jsonify({
                'message': 'Duplicate resume detected',
                'status': 'duplicate',
                'data': existing
            }), 200
    
    settings = get_scoring_settings()
    candidate_record = {
        'id': identifier,
        'filename': filename,
        'filepath': filepath,
        'status': 'uploaded',
        'retry_count': 0,
        'synced': False,
        'uploaded_at': datetime.now().isoformat(),
        'raw_text': text[:1000],
        'scoring_version': settings.get('version', '1.0')
    }
    
    save_candidate(identifier, candidate_record)
    
    parsed_data, error = parse_resume_with_openai(text)
    
    if error:
        candidate_record['retry_count'] += 1
        candidate_record['error'] = error
    elif parsed_data and validate_parsed_data(parsed_data):
        address_str = ""
        if parsed_data.get('address'):
            addr = parsed_data['address']
            address_str = " ".join([str(addr.get(k, "")) for k in ['street', 'city', 'state', 'zip_code', 'country']]).strip()
        
        parsed_data['service_location'] = classify_location(address_str, parsed_data.get('mobile', ''))
        
        candidate_record['status'] = 'processed'
        candidate_record['parsed_data'] = parsed_data
        candidate_record['processed_at'] = datetime.now().isoformat()
        
        success, message = send_to_zoho_flow(parsed_data)
        if success:
            candidate_record['synced'] = True
            candidate_record['zoho_synced_at'] = datetime.now().isoformat()
        
        save_candidate(identifier, candidate_record)
        
        return jsonify({
            'status': 'processed',
            'parsed_data': parsed_data,
            'synced': success,
            'message': message
        }), 200
    else:
        candidate_record['retry_count'] += 1
        missing_fields = [f for f in ['name', 'email', 'tier_level', 'tier_score', 'qualify'] if f not in parsed_data or not parsed_data.get(f)]
        candidate_record['error'] = f'Missing required fields: {", ".join(missing_fields)}'
        
        if candidate_record['retry_count'] < 3:
            candidate_record['status'] = 'uploaded'
        else:
            candidate_record['status'] = 'failed'
        
        save_candidate(identifier, candidate_record)
        
        return jsonify({
            'status': candidate_record['status'],
            'error': candidate_record['error'],
            'retry_count': candidate_record['retry_count']
        }), 200

@app.route('/retry/<candidate_id>', methods=['POST'])
@require_auth
def retry_candidate(candidate_id):
    db = get_candidates_db()
    
    if candidate_id not in db:
        return jsonify({'error': 'Candidate not found'}), 404
    
    candidate = db[candidate_id]
    
    if candidate['status'] == 'processed':
        return jsonify({'message': 'Candidate already processed'}), 200
    
    text = candidate.get('raw_text', '')
    if not text:
        text = extract_text(candidate['filepath'], candidate['filename'].rsplit('.', 1)[1].lower())
    
    parsed_data, error = parse_resume_with_openai(text)
    
    if error:
        candidate['retry_count'] += 1
        candidate['error'] = error
    elif parsed_data and validate_parsed_data(parsed_data):
        address_str = ""
        if parsed_data.get('address'):
            addr = parsed_data['address']
            address_str = " ".join([str(addr.get(k, "")) for k in ['street', 'city', 'state', 'zip_code', 'country']]).strip()
        
        parsed_data['service_location'] = classify_location(address_str, parsed_data.get('mobile', ''))
        
        candidate['status'] = 'processed'
        candidate['parsed_data'] = parsed_data
        candidate['processed_at'] = datetime.now().isoformat()
        
        success, message = send_to_zoho_flow(parsed_data)
        if success:
            candidate['synced'] = True
            candidate['zoho_synced_at'] = datetime.now().isoformat()
        
        save_candidate(candidate_id, candidate)
        
        return jsonify({
            'status': 'processed',
            'parsed_data': parsed_data,
            'synced': success
        }), 200
    else:
        candidate['retry_count'] += 1
        missing_fields = [f for f in ['name', 'email', 'tier_level', 'tier_score', 'qualify'] if f not in parsed_data or not parsed_data.get(f)]
        candidate['error'] = f'Missing required fields: {", ".join(missing_fields)}'
    
    if candidate['retry_count'] < 3:
        candidate['status'] = 'uploaded'
    else:
        candidate['status'] = 'failed'
    
    save_candidate(candidate_id, candidate)
    
    return jsonify({
        'status': candidate['status'],
        'error': candidate['error'],
        'retry_count': candidate['retry_count']
    }), 200

@app.route('/candidates', methods=['GET'])
@require_auth
def get_candidates():
    db = get_candidates_db()
    status_filter = request.args.get('status')
    
    if status_filter:
        filtered = {k: v for k, v in db.items() if v['status'] == status_filter}
        return jsonify(filtered), 200
    
    return jsonify(db), 200

@app.route('/settings', methods=['GET'])
@require_auth
def get_settings():
    return jsonify(get_scoring_settings()), 200

@app.route('/settings', methods=['POST'])
@require_auth
def update_settings():
    new_settings = request.json
    save_scoring_settings(new_settings)
    return jsonify({'message': 'Settings updated successfully', 'settings': new_settings}), 200

@app.route('/settings/page')
@require_auth
def settings_page():
    return render_template_string(open('settings.html').read())

@app.route('/dashboard')
@require_auth
def dashboard():
    return render_template_string(open('dashboard.html').read())

@app.route('/export/csv')
@require_auth
def export_csv():
    db = get_candidates_db()
    candidates = list(db.values())
    
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Headers
    headers = [
        'Name', 'Email', 'Primary Language', 'Other Languages', 'Service Location',
        'Mobile', 'Remote Experience', 'Tier Level', 'Tier Score', 'Education',
        'Qualify', 'Role Relevance', 'Training Needed', 'Processing Notes',
        'Status', 'Uploaded At', 'Processed At', 'Synced', 'Address'
    ]
    writer.writerow(headers)
    
    # Data rows
    for candidate in candidates:
        data = candidate.get('parsed_data', {})
        address_parts = []
        if data.get('address'):
            addr = data['address']
            address_parts = [
                addr.get('street', ''),
                addr.get('city', ''),
                addr.get('state', ''),
                addr.get('zip_code', ''),
                addr.get('country', '')
            ]
        
        row = [
            data.get('name', ''),
            data.get('email', candidate.get('id', '')),
            data.get('primary_language', ''),
            '; '.join(data.get('other_spoken_languages', [])) if isinstance(data.get('other_spoken_languages'), list) else '',
            data.get('service_location', ''),
            data.get('mobile', ''),
            'Yes' if data.get('remote_experience') else 'No',
            data.get('tier_level', ''),
            data.get('tier_score', 0),
            data.get('education', ''),
            data.get('qualify', ''),
            data.get('role_relevance', ''),
            'Yes' if data.get('training_needed') else 'No',
            data.get('processing_notes', ''),
            candidate.get('status', ''),
            candidate.get('uploaded_at', ''),
            candidate.get('processed_at', ''),
            'Yes' if candidate.get('synced') else 'No',
            ', '.join(filter(None, address_parts))
        ]
        writer.writerow(row)
    
    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename=candidates_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'}
    )

@app.route('/export/excel')
@require_auth
def export_excel():
    db = get_candidates_db()
    candidates = list(db.values())
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Candidates"
    
    # Headers
    headers = [
        'Name', 'Email', 'Primary Language', 'Other Languages', 'Service Location',
        'Mobile', 'Remote Experience', 'Tier Level', 'Tier Score', 'Education',
        'Qualify', 'Role Relevance', 'Training Needed', 'Processing Notes',
        'Status', 'Uploaded At', 'Processed At', 'Synced', 'Address'
    ]
    ws.append(headers)
    
    # Data rows
    for candidate in candidates:
        data = candidate.get('parsed_data', {})
        address_parts = []
        if data.get('address'):
            addr = data['address']
            address_parts = [
                addr.get('street', ''),
                addr.get('city', ''),
                addr.get('state', ''),
                addr.get('zip_code', ''),
                addr.get('country', '')
            ]
        
        row = [
            data.get('name', ''),
            data.get('email', candidate.get('id', '')),
            data.get('primary_language', ''),
            '; '.join(data.get('other_spoken_languages', [])) if isinstance(data.get('other_spoken_languages'), list) else '',
            data.get('service_location', ''),
            data.get('mobile', ''),
            'Yes' if data.get('remote_experience') else 'No',
            data.get('tier_level', ''),
            data.get('tier_score', 0),
            data.get('education', ''),
            data.get('qualify', ''),
            data.get('role_relevance', ''),
            'Yes' if data.get('training_needed') else 'No',
            data.get('processing_notes', ''),
            candidate.get('status', ''),
            candidate.get('uploaded_at', ''),
            candidate.get('processed_at', ''),
            'Yes' if candidate.get('synced') else 'No',
            ', '.join(filter(None, address_parts))
        ]
        ws.append(row)
    
    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    return Response(
        output.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=candidates_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'}
    )

@app.route('/')
def index():
    with open('index.html', 'r') as f:
        return render_template_string(f.read())

if __name__ == '__main__':
    # Use environment variable for debug mode, default to False for security
    debug_mode = os.getenv('FLASK_DEBUG', 'False').lower() == 'true'
    # Bind to localhost only for security, unless explicitly configured
    host = os.getenv('FLASK_HOST', '127.0.0.1')
    app.run(debug=debug_mode, host=host, port=5001)